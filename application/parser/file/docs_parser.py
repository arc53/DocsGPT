"""Docs parser.

Contains parsers for docx, pdf files.

"""

from pathlib import Path
from typing import Dict, List, Any, Union, Optional
from base64 import b64encode
from application.parser.file.base_parser import BaseParser
from application.core.settings import settings
from docx.enum.shape import WD_INLINE_SHAPE_TYPE
import requests
import logging
import sys
from docx import Document
import base64
from PIL import Image
import io

logger = logging.getLogger(__name__)


class PDFParser(BaseParser):
    """PDF parser."""

    def _init_parser(self) -> Dict:
        """Init parser."""
        return {}

    def parse_file(self, file: Path, errors: str = "ignore") -> str:
        """Parse file."""
        if settings.PARSE_PDF_AS_IMAGE:
            doc2md_service = "https://llm.arc53.com/doc2md"
            # alternatively you can use local vision capable LLM
            with open(file, "rb") as file_loaded:
                files = {"file": file_loaded}
                response = requests.post(doc2md_service, files=files)
                data = response.json()["markdown"]
            return data

        try:
            import PyPDF2
        except ImportError:
            raise ValueError("PyPDF2 is required to read PDF files.")
        text_list = []
        with open(file, "rb") as fp:
            # Create a PDF object
            pdf = PyPDF2.PdfReader(fp)

            # Get the number of pages in the PDF document
            num_pages = len(pdf.pages)

            # Iterate over every page
            for page in range(num_pages):
                # Extract the text from the page
                page_text = pdf.pages[page].extract_text()
                text_list.append(page_text)
        text = "\n".join(text_list)

        return text


class DocxParser(BaseParser):
    """Docx parser to extract text, tables, and images."""

    def _init_parser(self) -> Dict:
        """Init parser."""
        return {}

    def parse_file(self, file: Path, errors: str = "ignore") -> Dict[str, Any]:
        """Parse file and extract text, tables, and images."""
        document = Document(file)
        text_content = []
        tables = []
        images = []

        # Extract text from paragraphs
        for para in document.paragraphs:
            if para.text.strip():
                text_content.append(para.text.strip())

        # Extract tables
        for table in document.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(row_data)

            # Flatten table into a string
            flat_table = "\n".join([" | ".join(row) for row in table_data])
            tables.append(flat_table)

        # Extract images
        for shape in document.inline_shapes:
            if shape.type == WD_INLINE_SHAPE_TYPE.PICTURE:
                image_data = self.extract_image_from_shape(shape, document=document)
                if image_data:
                    images.append(image_data)

        return {"text": "\n".join(text_content), "tables": tables, "images": images}

    def extract_image_from_shape(self, shape, document) -> Optional[Dict[str, str]]:
        """Extract image from an inline shape and encode it in Base64."""
        try:
            rId = shape._inline.graphic.graphicData.pic.blipFill.blip.embed
            image_part = document.part.related_parts[rId]
            image_data = image_part.blob
            image_filename = image_part.filename or f"image_{rId}.png"

            # Validate the image
            if not self.is_valid_image(image_data):
                print(f"Invalid image: {image_filename}")
                return None
            resized_image = self.resize_image(image_data)
            image_base64 = base64.b64encode(resized_image).decode("utf-8")
            return {"filename": image_filename, "image_base64": image_base64}
        except Exception as e:
            print(f"Error extracting image: {e}")
            return None

    def is_valid_image(self, image_data: bytes) -> bool:
        """Validate image data."""
        try:
            image = Image.open(io.BytesIO(image_data))
            image.verify()
            return True
        except Exception:
            return False

    def resize_image(self, image_data: bytes, width=200) -> bytes:
        """Resize image to a specified maximum width, maintaining aspect ratio."""
        image = Image.open(io.BytesIO(image_data))
        ratio = width / float(image.width)
        height = int(image.height * ratio)
        img = image.resize((width, height), Image.Resampling.LANCZOS)

        # Save the resized image back to bytes
        output = io.BytesIO()
        # Use original format if possible, else default to PNG
        img_format = image.format if image.format else "PNG"
        img.save(output, format=img_format)
        return output.getvalue()
